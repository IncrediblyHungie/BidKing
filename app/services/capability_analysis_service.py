"""
Capability Statement Analysis Service

Uses Claude API to analyze capability statement PDFs and extract structured information
for personalized opportunity scoring.
"""

import json
import logging
from datetime import datetime
from typing import Optional, Dict, Any

import anthropic

from app.config import settings

logger = logging.getLogger(__name__)

# The prompt for Claude to extract structured data from capability statements
# NOTE: Literal braces in the JSON template are doubled ({{ and }}) to escape them from .format()
CAPABILITY_EXTRACTION_PROMPT = """You are an expert federal contracting analyst specializing in capability statement analysis. Analyze this company capability statement document and extract key information for matching against federal contract opportunities.

<document>
{document_text}
</document>

Extract the following information and return it as JSON. If information is not found, use null for that field.

{{
  "company_name": "The company name from the capability statement",
  "core_competencies": [
    "List of core competencies/services the company offers",
    "Focus on specific technical capabilities",
    "Include areas like 'Cloud Migration', 'Cybersecurity', 'Data Analytics', etc."
  ],
  "differentiators": [
    "What makes this company unique or competitive",
    "Special qualifications, approaches, or capabilities",
    "Past performance highlights"
  ],
  "keywords": [
    "Specific technical keywords for opportunity matching",
    "Technologies, methodologies, certifications mentioned",
    "Industry-specific terms (e.g., 'FedRAMP', 'DevSecOps', 'Zero Trust')",
    "Extract 15-30 highly relevant keywords"
  ],
  "target_naics_codes": [
    "Any NAICS codes explicitly mentioned in the document",
    "Format as 6-digit strings like '541511', '541512'"
  ],
  "target_agencies": [
    "Federal agencies the company has worked with or targets",
    "E.g., 'Department of Defense', 'HHS', 'DHS', 'NASA'"
  ],
  "technologies": [
    "Specific technologies, platforms, and tools mentioned",
    "E.g., 'AWS', 'Azure', 'Kubernetes', 'ServiceNow', 'Salesforce'"
  ],
  "certifications_mentioned": [
    "Any certifications mentioned (company or personnel)",
    "E.g., 'CMMI Level 3', 'ISO 27001', 'FedRAMP', 'PMP', 'CISSP'"
  ],
  "past_performance_summary": "Brief summary of notable past performance or contract experience mentioned",
  "contact_info": {{
    "email": "Contact email if found",
    "phone": "Contact phone if found",
    "website": "Company website if found",
    "address": "Address if found"
  }}
}}

Important guidelines:
1. For keywords, extract terms that would match federal solicitation requirements
2. Focus on technical capabilities, methodologies, and specific expertise areas
3. Include both broad categories (e.g., 'IT Services') and specific skills (e.g., 'Python Development')
4. For NAICS codes, only include if explicitly mentioned with numbers
5. Prioritize keywords that indicate capability to perform federal work
6. Return ONLY valid JSON, no other text"""


def analyze_capability_statement(
    text_content: str,
    file_name: str = "",
    max_chars: int = 100000
) -> Dict[str, Any]:
    """
    Use Claude to analyze capability statement text and extract structured information.

    Args:
        text_content: The extracted text from the PDF
        file_name: Name of the file for logging
        max_chars: Maximum characters to send to Claude (to manage costs)

    Returns:
        Dictionary with extracted information or error details
    """
    if not settings.anthropic_api_key:
        logger.error("Anthropic API key not configured")
        return {
            "error": "API key not configured",
            "status": "failed"
        }

    if not text_content or len(text_content.strip()) < 100:
        logger.info(f"Skipping {file_name}: insufficient text content")
        return {
            "error": "Insufficient text content",
            "status": "skipped"
        }

    # Truncate if too long (to manage API costs)
    if len(text_content) > max_chars:
        text_content = text_content[:max_chars] + "\n\n[Document truncated for analysis...]"

    try:
        client = anthropic.Anthropic(api_key=settings.anthropic_api_key)

        # Build the prompt
        prompt = CAPABILITY_EXTRACTION_PROMPT.format(document_text=text_content)

        # Call Claude API
        logger.info(f"Calling Claude API for capability statement: {file_name}...")
        message = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=4096,
            messages=[
                {
                    "role": "user",
                    "content": prompt
                }
            ]
        )
        logger.info(f"Claude API response received, stop_reason: {message.stop_reason}")

        # Extract the text response
        response_text = ""
        if message.content and len(message.content) > 0:
            block = message.content[0]
            logger.info(f"Content block type: {type(block).__name__}")
            if hasattr(block, 'text'):
                response_text = block.text
            else:
                logger.error(f"Block has no text attribute: {block}")

        if not response_text:
            return {
                "error": "No text response from Claude",
                "status": "failed"
            }

        # Parse JSON from response
        # Claude might wrap it in markdown code blocks
        json_text = response_text.strip()
        if json_text.startswith("```json"):
            json_text = json_text[7:]
        elif json_text.startswith("```"):
            json_text = json_text[3:]
        if json_text.endswith("```"):
            json_text = json_text[:-3]
        json_text = json_text.strip()

        # Find the actual JSON object - look for first { and last }
        first_brace = json_text.find('{')
        last_brace = json_text.rfind('}')
        if first_brace != -1 and last_brace != -1:
            json_text = json_text[first_brace:last_brace + 1]

        try:
            analysis_data = json.loads(json_text)
            analysis_data["status"] = "analyzed"
            analysis_data["model"] = "claude-sonnet-4-20250514"
            analysis_data["analyzed_at"] = datetime.utcnow().isoformat()
            return analysis_data
        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse Claude response as JSON: {e}")
            logger.error(f"JSON text was: {json_text[:500]}")
            return {
                "error": f"JSON parse error at position {e.pos}: {e.msg}",
                "raw_response": response_text[:500],
                "status": "failed"
            }
        except Exception as e:
            logger.error(f"Unexpected error parsing Claude response: {type(e).__name__}: {e}")
            return {
                "error": f"Parse error: {type(e).__name__}: {str(e)[:200]}",
                "raw_response": response_text[:500] if response_text else "No response",
                "status": "failed"
            }

    except anthropic.APIError as e:
        logger.error(f"Anthropic API error: {e}")
        return {
            "error": f"API error: {str(e)}",
            "status": "failed"
        }
    except Exception as e:
        logger.error(f"Error analyzing capability statement: {e}")
        return {
            "error": str(e),
            "status": "failed"
        }


def extract_text_from_pdf(pdf_bytes: bytes, file_name: str = "") -> Optional[str]:
    """
    Extract text content from a PDF file.

    Args:
        pdf_bytes: The PDF file as bytes
        file_name: Name of the file for logging

    Returns:
        Extracted text or None if extraction fails
    """
    try:
        import pypdf
        from io import BytesIO

        reader = pypdf.PdfReader(BytesIO(pdf_bytes))
        text_parts = []

        for page_num, page in enumerate(reader.pages):
            try:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            except Exception as e:
                logger.warning(f"Failed to extract text from page {page_num + 1} of {file_name}: {e}")
                continue

        full_text = "\n\n".join(text_parts)
        logger.info(f"Extracted {len(full_text)} characters from {file_name} ({len(reader.pages)} pages)")
        return full_text if full_text.strip() else None

    except Exception as e:
        logger.error(f"Failed to extract text from PDF {file_name}: {e}")
        return None


def extract_text_from_docx(docx_bytes: bytes, file_name: str = "") -> Optional[str]:
    """
    Extract text content from a DOCX file.

    Args:
        docx_bytes: The DOCX file as bytes
        file_name: Name of the file for logging

    Returns:
        Extracted text or None if extraction fails
    """
    try:
        from docx import Document
        from io import BytesIO

        doc = Document(BytesIO(docx_bytes))
        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)

        full_text = "\n\n".join(text_parts)
        logger.info(f"Extracted {len(full_text)} characters from {file_name}")
        return full_text if full_text.strip() else None

    except Exception as e:
        logger.error(f"Failed to extract text from DOCX {file_name}: {e}")
        return None
