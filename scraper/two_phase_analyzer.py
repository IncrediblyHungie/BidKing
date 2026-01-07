#!/usr/bin/env python3
"""
Two-Phase AI Analysis System for SAM.gov Opportunities

Phase 1: Extract text from ALL documents (fast, parallelizable)
  - Detects file type by magic bytes (not extension)
  - Stores extracted text in database
  - Can be run independently

Phase 2: AI Analysis per OPPORTUNITY (not per document)
  - Groups all documents for each opportunity
  - Concatenates all text together
  - Makes ONE LLM call per opportunity
  - Better context = better extraction

This approach:
- Fixes file extension detection issues
- Reduces LLM calls significantly
- Gives LLM full context across all documents
"""

import asyncio
import json
import logging
import re
import sqlite3
import subprocess
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Optional, Dict, List, Any
from concurrent.futures import ThreadPoolExecutor

import httpx

# PDF extraction
try:
    import pypdf
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

# PyMuPDF (fitz) - FASTEST PDF extraction
try:
    import fitz  # pymupdf
    HAS_FITZ = True
except ImportError:
    HAS_FITZ = False

# DOCX extraction
try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# XLSX extraction
try:
    import openpyxl
    HAS_XLSX = True
except ImportError:
    HAS_XLSX = False

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('two_phase_analyzer.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)


# =============================================================================
# Magic Byte Detection
# =============================================================================

MAGIC_BYTES = {
    b'%PDF': 'pdf',
    b'PK\x03\x04': 'zip',  # DOCX, XLSX, PPTX are all ZIP-based
    b'\xd0\xcf\x11\xe0': 'ole',  # Old Office formats (DOC, XLS, PPT)
    b'{\rtf': 'rtf',
}

def detect_file_type(file_path: Path) -> str:
    """Detect file type by magic bytes, not extension."""
    try:
        with open(file_path, 'rb') as f:
            header = f.read(16)

        # Check magic bytes
        if header.startswith(b'%PDF'):
            return 'pdf'

        if header.startswith(b'PK\x03\x04'):
            # ZIP-based format - need to check internal structure
            # DOCX has word/document.xml, XLSX has xl/workbook.xml
            import zipfile
            try:
                with zipfile.ZipFile(file_path) as zf:
                    names = zf.namelist()
                    if any('word/' in n for n in names):
                        return 'docx'
                    elif any('xl/' in n for n in names):
                        return 'xlsx'
                    elif any('ppt/' in n for n in names):
                        return 'pptx'
                    else:
                        return 'zip'
            except:
                return 'zip'

        if header.startswith(b'\xd0\xcf\x11\xe0'):
            # OLE format - old Office
            # Try to determine if DOC or XLS by extension as fallback
            ext = file_path.suffix.lower()
            if ext == '.xls':
                return 'xls'
            return 'doc'  # Default to DOC for OLE

        if header.startswith(b'{\\rtf'):
            return 'rtf'

        # Check if it's plain text (mostly printable ASCII)
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                sample = f.read(1000)
                if sample and all(c.isprintable() or c in '\n\r\t' for c in sample):
                    return 'txt'
        except:
            pass

        # Fallback to extension
        ext = file_path.suffix.lower()
        if ext in ['.pdf', '.docx', '.doc', '.xlsx', '.xls', '.txt', '.rtf']:
            return ext[1:]

        return 'unknown'

    except Exception as e:
        logger.warning(f"Could not detect file type for {file_path}: {e}")
        return 'unknown'




# =============================================================================
# Configuration
# =============================================================================

@dataclass
class AnalysisConfig:
    """Configuration for the analysis system."""

    ollama_host: str = "http://localhost:11434"
    model: str = "qwen2.5:14b-instruct"
    fallback_model: str = "qwen2.5:7b-instruct"

    scraper_db: str = "bidking_sam.db"
    pdf_dir: str = "pdfs"

    # For multi-document opportunities, we can send more text
    max_text_per_doc: int = 20000  # Max chars per document
    max_total_text: int = 60000   # Max total chars for all docs combined

    llm_timeout: int = 300  # 5 minutes per analysis


# =============================================================================
# Database Schema Extension
# =============================================================================

def init_tables(db_path: str):
    """Initialize tables for two-phase analysis."""

    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()

    # Extended attachments table for text extraction
    cursor.execute("PRAGMA table_info(attachments)")
    existing_cols = {row[1] for row in cursor.fetchall()}

    if 'text_extracted' not in existing_cols:
        cursor.execute("ALTER TABLE attachments ADD COLUMN text_extracted INTEGER DEFAULT 0")
    if 'extracted_text' not in existing_cols:
        cursor.execute("ALTER TABLE attachments ADD COLUMN extracted_text TEXT")
    if 'detected_type' not in existing_cols:
        cursor.execute("ALTER TABLE attachments ADD COLUMN detected_type TEXT")
    if 'extraction_error' not in existing_cols:
        cursor.execute("ALTER TABLE attachments ADD COLUMN extraction_error TEXT")

    # Per-OPPORTUNITY AI analysis (not per-document)
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS opportunity_analysis (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            opportunity_id TEXT UNIQUE NOT NULL,

            -- Analysis status
            status TEXT DEFAULT 'pending',  -- pending, completed, failed, no_docs
            error_message TEXT,

            -- Combined document info
            num_documents INTEGER,
            total_text_length INTEGER,
            source_documents TEXT,  -- JSON array of filenames

            -- AI Analysis JSON (matches BidKing schema)
            ai_summary JSON,

            -- Metadata
            model_used TEXT,
            analysis_duration_seconds REAL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            analyzed_at TIMESTAMP,

            FOREIGN KEY (opportunity_id) REFERENCES opportunities(opportunity_id)
        )
    """)

    cursor.execute("""
        CREATE INDEX IF NOT EXISTS idx_opp_analysis_status
        ON opportunity_analysis(status)
    """)

    conn.commit()
    conn.close()
    logger.info(f"Initialized two-phase analysis tables in {db_path}")


# =============================================================================
# Phase 1: Text Extraction
# =============================================================================

class TextExtractor:
    """Phase 1: Extract text from all downloaded documents."""

    def __init__(self, config: AnalysisConfig):
        self.config = config
        self.pdf_dir = Path(config.pdf_dir)

    def extract_text(self, file_path: Path) -> Optional[str]:
        """Extract text from any document type using magic byte detection."""

        if not file_path.exists():
            return None

        file_type = detect_file_type(file_path)
        logger.debug(f"Detected {file_path.name} as {file_type}")

        text = None

        if file_type == 'pdf':
            text = self._extract_pdf(file_path)
        elif file_type == 'docx':
            text = self._extract_docx(file_path)
        elif file_type == 'doc':
            text = self._extract_doc(file_path)
        elif file_type == 'xlsx':
            text = self._extract_xlsx(file_path)
        elif file_type == 'xls':
            text = self._extract_xlsx(file_path)  # openpyxl handles both
        elif file_type == 'txt':
            text = self._extract_txt(file_path)
        elif file_type == 'rtf':
            text = self._extract_txt(file_path)
        else:
            logger.warning(f"Unsupported file type '{file_type}' for {file_path.name}")
            return None

        return text.strip() if text else None

    def _extract_pdf(self, path: Path) -> Optional[str]:
        """Extract text from PDF using pdfplumber (thread-safe) with pypdf fallback."""
        text = ""

        # NOTE: PyMuPDF (fitz) is faster but NOT thread-safe with ThreadPoolExecutor
        # Using pdfplumber as primary which is thread-safe
        if HAS_PDFPLUMBER:
            try:
                with pdfplumber.open(path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n\n"
                if text.strip():
                    return text.strip()
            except Exception as e:
                logger.debug(f"pdfplumber failed: {e}")

        # Final fallback to pypdf
        if HAS_PYPDF:
            try:
                reader = pypdf.PdfReader(path)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
                if text.strip():
                    return text.strip()
            except Exception as e:
                logger.debug(f"pypdf failed: {e}")

        return None

    def _extract_docx(self, path: Path) -> Optional[str]:
        """Extract text from DOCX."""
        if not HAS_DOCX:
            return None

        try:
            doc = DocxDocument(path)
            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            return "\n\n".join(text_parts)
        except Exception as e:
            logger.debug(f"DOCX extraction failed: {e}")
            return None

    def _extract_doc(self, path: Path) -> Optional[str]:
        """Extract text from old DOC format."""
        try:
            result = subprocess.run(
                ['antiword', str(path)],
                capture_output=True, text=True, timeout=30
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout.strip()
        except:
            pass

        # Fallback: extract printable ASCII
        try:
            with open(path, 'rb') as f:
                content = f.read()
                text = ''.join(chr(b) for b in content if 32 <= b < 127 or b in (10, 13, 9))
                text = ' '.join(text.split())
                if len(text) > 100:
                    return text
        except:
            pass

        return None

    def _extract_xlsx(self, path: Path) -> Optional[str]:
        """Extract text from XLSX/XLS."""
        if not HAS_XLSX:
            return None

        try:
            wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
            text_parts = []

            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                text_parts.append(f"=== Sheet: {sheet_name} ===")

                for row in sheet.iter_rows():
                    row_values = []
                    for cell in row:
                        if cell.value is not None:
                            row_values.append(str(cell.value).strip())
                    if row_values:
                        text_parts.append(" | ".join(row_values))

            wb.close()
            return "\n".join(text_parts)
        except Exception as e:
            logger.debug(f"XLSX extraction failed: {e}")
            return None

    def _extract_txt(self, path: Path) -> Optional[str]:
        """Extract text from plain text file."""
        for encoding in ['utf-8', 'latin-1', 'cp1252']:
            try:
                with open(path, 'r', encoding=encoding) as f:
                    return f.read()
            except:
                continue
        return None

    def get_pending_documents(self, limit: int = 1000) -> List[Dict]:
        """Get downloaded documents that haven't had text extracted."""

        conn = sqlite3.connect(self.config.scraper_db)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()

        cursor.execute("""
            SELECT
                id, opportunity_id, filename, resource_id,
                pdf_local_path, mime_type
            FROM attachments
            WHERE pdf_downloaded = 1
            AND (text_extracted = 0 OR text_extracted IS NULL)
            AND pdf_local_path IS NOT NULL
            ORDER BY id
            LIMIT ?
        """, (limit,))

        results = [dict(row) for row in cursor.fetchall()]
        conn.close()
        return results

    def save_extracted_text(self, att_id: int, text: Optional[str],
                           detected_type: str, error: Optional[str] = None):
        """Save extracted text to database."""

        conn = sqlite3.connect(self.config.scraper_db)
        cursor = conn.cursor()

        cursor.execute("""
            UPDATE attachments
            SET text_extracted = 1,
                extracted_text = ?,
                detected_type = ?,
                extraction_error = ?
            WHERE id = ?
        """, (text, detected_type, error, att_id))

        conn.commit()
        conn.close()


    def _process_single_doc(self, doc: Dict) -> tuple:
        """Process a single document for text extraction."""
        file_path = Path(doc['pdf_local_path'])
        att_id = doc['id']

        if not file_path.exists():
            return (att_id, None, 'missing', 'File not found')

        try:
            file_type = detect_file_type(file_path)
            text = self.extract_text(file_path)

            if text and len(text) >= 50:
                return (att_id, text, file_type, None)
            else:
                return (att_id, None, file_type, 'No extractable text')
        except Exception as e:
            return (att_id, None, 'error', str(e))

    def run_extraction(self, limit: int = 10000, workers: int = 16):
        """Run Phase 1: Extract text from all pending documents using process pool.

        Uses ProcessPoolExecutor instead of ThreadPoolExecutor for PyMuPDF thread safety.
        Default 8 workers to avoid memory bloat.
        """
        from concurrent.futures import ProcessPoolExecutor, as_completed
        import os

        logger.info("=" * 60)
        logger.info("PHASE 1: Text Extraction (Parallel - ProcessPool)")
        logger.info(f"Using {workers} worker processes ({os.cpu_count()} CPUs available)")
        logger.info("=" * 60)

        docs = self.get_pending_documents(limit)
        logger.info(f"Found {len(docs)} documents to process")

        if not docs:
            logger.info("No pending documents for text extraction")
            return

        extracted = 0
        failed = 0
        processed = 0

        # Use ThreadPoolExecutor for parallel extraction
        # Note: PyMuPDF (fitz) has threading issues - falls back to pdfplumber/pypdf which are thread-safe
        with ThreadPoolExecutor(max_workers=workers) as executor:
            # Submit all tasks
            future_to_doc = {executor.submit(self._process_single_doc, doc): doc for doc in docs}

            # Process results as they complete
            for future in as_completed(future_to_doc):
                processed += 1

                try:
                    att_id, text, file_type, error = future.result()
                    self.save_extracted_text(att_id, text, file_type, error)

                    if text:
                        extracted += 1
                    else:
                        failed += 1
                except Exception as e:
                    failed += 1
                    logger.error(f"Error processing document: {e}")

                # Log progress every 100 docs
                if processed % 100 == 0:
                    logger.info(f"Progress: {processed}/{len(docs)} ({extracted} extracted, {failed} failed)")

        logger.info(f"\nPhase 1 Complete: Extracted {extracted}, Failed {failed}")
        return

    def run_extraction_single_threaded(self, limit: int = 10000):
        """Run Phase 1: Extract text from all pending documents (single-threaded fallback)."""

        logger.info("=" * 60)
        logger.info("PHASE 1: Text Extraction (Single-threaded)")
        logger.info("=" * 60)

        docs = self.get_pending_documents(limit)
        logger.info(f"Found {len(docs)} documents to process")

        if not docs:
            logger.info("No pending documents for text extraction")
            return

        extracted = 0
        failed = 0

        for i, doc in enumerate(docs, 1):
            file_path = Path(doc['pdf_local_path'])
            filename = doc['filename'] or doc['resource_id']

            if not file_path.exists():
                self.save_extracted_text(doc['id'], None, 'missing', 'File not found')
                failed += 1
                continue

            try:
                file_type = detect_file_type(file_path)
                text = self.extract_text(file_path)

                if text and len(text) >= 50:
                    self.save_extracted_text(doc['id'], text, file_type)
                    extracted += 1
                    if i % 100 == 0:
                        logger.info(f"[{i}/{len(docs)}] Extracted: {extracted}, Failed: {failed}")
                else:
                    self.save_extracted_text(doc['id'], None, file_type, 'No extractable text')
                    failed += 1

            except Exception as e:
                self.save_extracted_text(doc['id'], None, 'error', str(e))
                failed += 1

        logger.info(f"\nPhase 1 Complete: Extracted {extracted}, Failed {failed}")


# =============================================================================
# Phase 2: Per-Opportunity AI Analysis
# =============================================================================

class OpportunityAnalyzer:
    """Phase 2: Analyze opportunities using combined text from all documents."""

    ANALYSIS_PROMPT = """You are an expert federal contracting analyst. Analyze these government solicitation documents and extract key information in a structured JSON format.

IMPORTANT: These documents are ALL from the SAME opportunity. Use information from ALL documents to build a complete picture. Information may be spread across multiple files.

Combined Document Text:
{document_text}

---

Extract the following information and return ONLY a valid JSON object (no markdown, no explanation):

{{
    "summary": "2-3 sentence plain English summary of what the government is looking for",

    "estimated_value": {{
        "low": <number or null>,
        "high": <number or null>,
        "basis": "Explanation of how you derived the estimate"
    }},

    "period_of_performance": "e.g., '1 base year + 4 option years' or null if not stated",

    "contract_type": "e.g., 'Firm Fixed Price', 'Time & Materials', 'IDIQ', or null",

    "clearance_required": "e.g., 'None', 'Public Trust', 'Secret', 'Top Secret', 'TS/SCI', or null",

    "labor_categories": [
        {{"title": "Role title", "quantity": <number or null>, "level": "Junior/Mid/Senior/Lead/SME or null"}}
    ],

    "technologies": ["List of specific technologies, platforms, tools, languages mentioned"],

    "certifications_required": ["List of required certifications like CMMI, ISO, FedRAMP, etc."],

    "location": "e.g., 'Remote', 'Washington, DC', or description of work location",

    "incumbent": "Current contractor name if this is a recompete, or null",

    "key_dates": {{
        "proposal_due": "YYYY-MM-DD or null",
        "questions_due": "YYYY-MM-DD or null"
    }},

    "evaluation_factors": ["List of how proposals will be evaluated"],

    "small_business_set_aside": "e.g., '8(a)', 'SDVOSB', 'HUBZone', 'Full and Open', or null",

    "naics_codes": ["List of NAICS codes mentioned"],

    "deliverables": ["Key deliverables or work products expected"]
}}

Return ONLY the JSON object, no other text."""

    def __init__(self, config: AnalysisConfig):
        self.config = config
        self.model = config.model

    async def check_ollama(self) -> bool:
        """Check if Ollama is available."""
        try:
            async with httpx.AsyncClient(timeout=10) as client:
                response = await client.get(f"{self.config.ollama_host}/api/tags")
                if response.status_code == 200:
                    models = [m['name'] for m in response.json().get('models', [])]
                    if any(self.model in m for m in models):
                        return True
                    if any(self.config.fallback_model in m for m in models):
                        self.model = self.config.fallback_model
                        return True
            return False
        except:
            return False

    def get_opportunities_for_analysis(self, limit: int = 1000) -> List[str]:
        """Get opportunity IDs that have extracted text but no analysis."""

        conn = sqlite3.connect(self.config.scraper_db)
        cursor = conn.cursor()

        cursor.execute("""
            SELECT DISTINCT a.opportunity_id
            FROM attachments a
            LEFT JOIN opportunity_analysis oa ON a.opportunity_id = oa.opportunity_id
            WHERE a.text_extracted = 1
            AND a.extracted_text IS NOT NULL
            AND LENGTH(a.extracted_text) > 50
            AND oa.id IS NULL
            LIMIT ?
        """, (limit,))

        results = [row[0] for row in cursor.fetchall()]
        conn.close()
        return results

    def _score_document_importance(self, filename: str, text: str) -> int:
        """Score document importance for prioritization.

        Higher scores = more important documents.
        SOW/PWS > RFP/RFQ > Technical > Pricing > Other
        """
        score = 0
        filename_lower = filename.lower()
        text_lower = text[:5000].lower()  # Check first 5K chars

        # Highest priority: Statement of Work / Performance Work Statement
        sow_terms = ['statement of work', 'sow', 'pws', 'performance work statement',
                     'scope of work', 'work statement']
        if any(term in filename_lower for term in sow_terms):
            score += 100
        if any(term in text_lower for term in sow_terms):
            score += 50

        # High priority: RFP/RFQ documents
        rfp_terms = ['rfp', 'rfq', 'rfi', 'request for proposal', 'request for quote',
                     'solicitation', 'combined synopsis']
        if any(term in filename_lower for term in rfp_terms):
            score += 80
        if any(term in text_lower for term in rfp_terms):
            score += 40

        # High priority: Pricing / CLIN documents
        price_terms = ['pricing', 'clin', 'cost', 'price schedule', 'b.', 'section b']
        if any(term in filename_lower for term in price_terms):
            score += 70
        if any(term in text_lower for term in ['price schedule', 'clin', 'contract line item']):
            score += 35

        # Medium priority: Technical requirements
        tech_terms = ['technical', 'requirement', 'specification', 'section c', 'section l']
        if any(term in filename_lower for term in tech_terms):
            score += 60
        if any(term in text_lower for term in ['technical requirements', 'minimum qualifications']):
            score += 30

        # Medium priority: Evaluation criteria
        eval_terms = ['evaluation', 'criteria', 'section m', 'factor']
        if any(term in filename_lower for term in eval_terms):
            score += 50
        if any(term in text_lower for term in ['evaluation factor', 'evaluation criteria']):
            score += 25

        # Lower priority: Attachments that are typically less useful
        low_value_terms = ['sf ', 'form ', 'clauses', 'far ', 'dfar', '52.2', 'attachment j']
        if any(term in filename_lower for term in low_value_terms):
            score -= 20

        # Bonus: Longer documents often have more substance
        if len(text) > 10000:
            score += 15
        elif len(text) > 5000:
            score += 10

        return score

    def get_combined_text_for_opportunity(self, opp_id: str) -> tuple:
        """Get all extracted text for an opportunity, combined.

        For opportunities with many documents, prioritizes the most important
        documents (SOW, RFP, Pricing, Technical) to fit within context limits.
        """

        conn = sqlite3.connect(self.config.scraper_db)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()

        cursor.execute("""
            SELECT filename, extracted_text, detected_type
            FROM attachments
            WHERE opportunity_id = ?
            AND text_extracted = 1
            AND extracted_text IS NOT NULL
            AND LENGTH(extracted_text) > 50
            ORDER BY id
        """, (opp_id,))

        rows = cursor.fetchall()
        conn.close()

        if not rows:
            return None, [], 0

        total_docs = len(rows)

        # Score and sort documents by importance
        scored_docs = []
        for row in rows:
            filename = row['filename'] or 'Unknown'
            text = row['extracted_text']
            score = self._score_document_importance(filename, text)
            scored_docs.append({
                'filename': filename,
                'text': text,
                'score': score,
                'length': len(text)
            })

        # Sort by importance score (highest first)
        scored_docs.sort(key=lambda x: x['score'], reverse=True)

        # If many documents, log which ones we're prioritizing
        if len(scored_docs) > 10:
            top_5 = [f"{d['filename']} (score:{d['score']})" for d in scored_docs[:5]]
            logger.debug(f"Opportunity has {len(scored_docs)} docs, prioritizing: {top_5}")

        combined_parts = []
        filenames = []
        total_len = 0

        for doc in scored_docs:
            filename = doc['filename']
            text = doc['text']

            # Truncate individual docs if too long
            if len(text) > self.config.max_text_per_doc:
                text = text[:self.config.max_text_per_doc] + "\n[...truncated...]"

            # Check if adding this doc would exceed our limit
            new_len = total_len + len(text) + 100  # +100 for headers

            if new_len > self.config.max_total_text and combined_parts:
                # We have enough content, stop adding
                remaining = len(scored_docs) - len(filenames)
                if remaining > 0:
                    combined_parts.append(
                        f"\n[...{remaining} additional lower-priority documents not included...]"
                    )
                break

            combined_parts.append(f"\n{'='*40}\nDOCUMENT: {filename}\n{'='*40}\n{text}")
            filenames.append(filename)
            total_len += len(text)

        combined_text = "\n".join(combined_parts)

        # Final safety truncation if still too long
        if len(combined_text) > self.config.max_total_text:
            half = self.config.max_total_text // 2
            combined_text = combined_text[:half] + "\n\n[...content truncated...]\n\n" + combined_text[-half:]

        return combined_text, filenames, total_docs

    async def analyze_opportunity(self, opp_id: str) -> bool:
        """Analyze a single opportunity using all its documents."""

        combined_text, filenames, num_docs = self.get_combined_text_for_opportunity(opp_id)

        if not combined_text:
            self._save_result(opp_id, 'no_docs', None, 0, [], 0, 'No extractable text')
            return False

        logger.info(f"Analyzing {opp_id[:8]}... ({num_docs} docs, {len(combined_text):,} chars)")

        prompt = self.ANALYSIS_PROMPT.format(document_text=combined_text)
        start_time = datetime.now()

        try:
            async with httpx.AsyncClient(timeout=self.config.llm_timeout) as client:
                response = await client.post(
                    f"{self.config.ollama_host}/api/generate",
                    json={
                        "model": self.model,
                        "prompt": prompt,
                        "stream": False,
                        "options": {
                            "temperature": 0.1,
                            "num_predict": 4096,
                            "top_p": 0.9,
                        }
                    }
                )

                if response.status_code != 200:
                    duration = (datetime.now() - start_time).total_seconds()
                    self._save_result(opp_id, 'failed', None, num_docs, filenames,
                                     duration, f"Ollama error: {response.status_code}")
                    return False

                data = response.json()
                response_text = data.get('response', '')

                analysis = self._extract_json(response_text)
                duration = (datetime.now() - start_time).total_seconds()

                if analysis:
                    self._save_result(opp_id, 'completed', analysis, num_docs,
                                     filenames, duration)
                    logger.info(f"✅ Completed {opp_id[:8]} in {duration:.1f}s")
                    return True
                else:
                    self._save_result(opp_id, 'failed', None, num_docs, filenames,
                                     duration, 'Could not parse JSON')
                    return False

        except asyncio.TimeoutError:
            duration = (datetime.now() - start_time).total_seconds()
            self._save_result(opp_id, 'failed', None, num_docs, filenames,
                             duration, f"Timeout after {self.config.llm_timeout}s")
            return False
        except Exception as e:
            duration = (datetime.now() - start_time).total_seconds()
            self._save_result(opp_id, 'failed', None, num_docs, filenames,
                             duration, str(e))
            return False

    def _extract_json(self, text: str) -> Optional[Dict]:
        """Extract JSON from LLM response."""
        try:
            return json.loads(text)
        except:
            pass

        text = re.sub(r'```json\s*', '', text)
        text = re.sub(r'```\s*', '', text)

        try:
            start = text.find('{')
            end = text.rfind('}')
            if start != -1 and end != -1:
                return json.loads(text[start:end+1])
        except:
            pass

        return None

    def _save_result(self, opp_id: str, status: str, analysis: Optional[Dict],
                    num_docs: int, filenames: List[str], duration: float,
                    error: Optional[str] = None):
        """Save analysis result to database."""

        conn = sqlite3.connect(self.config.scraper_db)
        cursor = conn.cursor()

        cursor.execute("""
            INSERT OR REPLACE INTO opportunity_analysis (
                opportunity_id, status, error_message,
                num_documents, total_text_length, source_documents,
                ai_summary, model_used, analysis_duration_seconds, analyzed_at
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            opp_id,
            status,
            error,
            num_docs,
            0,  # Could compute if needed
            json.dumps(filenames),
            json.dumps(analysis) if analysis else None,
            self.model,
            duration,
            datetime.now().isoformat() if status == 'completed' else None
        ))

        conn.commit()
        conn.close()

    async def run_analysis(self, limit: int = 1000):
        """Run Phase 2: Analyze all pending opportunities."""

        logger.info("=" * 60)
        logger.info("PHASE 2: Per-Opportunity AI Analysis")
        logger.info("=" * 60)

        if not await self.check_ollama():
            logger.error("Ollama not available!")
            return

        logger.info(f"Using model: {self.model}")

        opp_ids = self.get_opportunities_for_analysis(limit)
        logger.info(f"Found {len(opp_ids)} opportunities to analyze")

        if not opp_ids:
            logger.info("No pending opportunities")
            return

        completed = 0
        failed = 0

        for i, opp_id in enumerate(opp_ids, 1):
            logger.info(f"\n[{i}/{len(opp_ids)}]")

            try:
                success = await self.analyze_opportunity(opp_id)
                if success:
                    completed += 1
                else:
                    failed += 1
            except Exception as e:
                logger.error(f"Error: {e}")
                failed += 1

            await asyncio.sleep(1)  # Brief pause

        logger.info(f"\nPhase 2 Complete: {completed} analyzed, {failed} failed")


# =============================================================================
# Statistics
# =============================================================================

def get_stats(db_path: str = "bidking_sam.db") -> Dict:
    """Get statistics for both phases."""

    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()

    stats = {}

    # Phase 1 stats
    cursor.execute("SELECT COUNT(*) FROM attachments WHERE pdf_downloaded = 1")
    stats['downloaded'] = cursor.fetchone()[0]

    cursor.execute("SELECT COUNT(*) FROM attachments WHERE text_extracted = 1 AND extracted_text IS NOT NULL")
    stats['text_extracted'] = cursor.fetchone()[0]

    cursor.execute("""
        SELECT COUNT(*) FROM attachments
        WHERE pdf_downloaded = 1
        AND (text_extracted = 0 OR text_extracted IS NULL)
    """)
    stats['extraction_pending'] = cursor.fetchone()[0]

    # Phase 2 stats
    cursor.execute("SELECT COUNT(DISTINCT opportunity_id) FROM attachments WHERE text_extracted = 1")
    stats['opportunities_with_text'] = cursor.fetchone()[0]

    cursor.execute("SELECT COUNT(*) FROM opportunity_analysis WHERE status = 'completed'")
    stats['analyzed'] = cursor.fetchone()[0]

    cursor.execute("SELECT COUNT(*) FROM opportunity_analysis WHERE status = 'failed'")
    stats['failed'] = cursor.fetchone()[0]

    cursor.execute("""
        SELECT COUNT(DISTINCT a.opportunity_id)
        FROM attachments a
        LEFT JOIN opportunity_analysis oa ON a.opportunity_id = oa.opportunity_id
        WHERE a.text_extracted = 1 AND a.extracted_text IS NOT NULL
        AND oa.id IS NULL
    """)
    stats['analysis_pending'] = cursor.fetchone()[0]

    conn.close()
    return stats


def print_stats(db_path: str = "bidking_sam.db"):
    """Print nicely formatted stats."""

    stats = get_stats(db_path)

    print("\n" + "=" * 50)
    print("Two-Phase Analysis Statistics")
    print("=" * 50)
    print("\nPHASE 1: Text Extraction")
    print(f"  Downloaded documents:    {stats['downloaded']:,}")
    print(f"  Text extracted:          {stats['text_extracted']:,}")
    print(f"  Pending extraction:      {stats['extraction_pending']:,}")
    print("\nPHASE 2: AI Analysis (per opportunity)")
    print(f"  Opportunities with text: {stats['opportunities_with_text']:,}")
    print(f"  Analyzed:                {stats['analyzed']:,}")
    print(f"  Failed:                  {stats['failed']:,}")
    print(f"  Pending analysis:        {stats['analysis_pending']:,}")
    print("=" * 50)


# =============================================================================
# Export for BidKing
# =============================================================================

def export_for_bidking(db_path: str = "bidking_sam.db",
                       output_file: str = "bidking_ai_import.json") -> int:
    """Export per-opportunity analysis results for BidKing."""

    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()

    cursor.execute("""
        SELECT
            o.opportunity_id,
            o.solicitation_number,
            o.title,
            o.agency_name,
            o.response_deadline,
            oa.ai_summary,
            oa.num_documents,
            oa.source_documents,
            oa.model_used,
            oa.analyzed_at
        FROM opportunity_analysis oa
        JOIN opportunities o ON oa.opportunity_id = o.opportunity_id
        WHERE oa.status = 'completed' AND oa.ai_summary IS NOT NULL
    """)

    export_data = []

    for row in cursor:
        ai_summary = json.loads(row['ai_summary']) if row['ai_summary'] else {}
        source_docs = json.loads(row['source_documents']) if row['source_documents'] else []

        export_data.append({
            "opportunity_id": row['opportunity_id'],
            "solicitation_number": row['solicitation_number'],
            "title": row['title'],
            "agency_name": row['agency_name'],
            "response_deadline": row['response_deadline'],
            "ai_summary": ai_summary,
            "source_documents": source_docs,
            "analysis_metadata": {
                "model_used": row['model_used'],
                "analyzed_at": row['analyzed_at'],
                "num_documents_analyzed": row['num_documents']
            }
        })

    conn.close()

    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump({
            "exported_at": datetime.now().isoformat(),
            "total_opportunities": len(export_data),
            "opportunities": export_data
        }, f, indent=2)

    logger.info(f"Exported {len(export_data)} opportunities to {output_file}")
    return len(export_data)


# =============================================================================
# CLI
# =============================================================================

async def main():
    import argparse

    parser = argparse.ArgumentParser(description="Two-Phase AI Analysis for SAM.gov")
    parser.add_argument("--db", default="bidking_sam.db", help="Database path")
    parser.add_argument("--phase1", action="store_true", help="Run Phase 1: Text extraction")
    parser.add_argument("--phase2", action="store_true", help="Run Phase 2: AI analysis")
    parser.add_argument("--all", action="store_true", help="Run both phases")
    parser.add_argument("--limit", type=int, default=10000, help="Max items to process")
    parser.add_argument("--model", default="qwen2.5:14b-instruct", help="Ollama model")
    parser.add_argument("--export", action="store_true", help="Export results for BidKing")
    parser.add_argument("--stats", action="store_true", help="Show statistics")

    args = parser.parse_args()

    # Initialize tables
    init_tables(args.db)

    if args.stats:
        print_stats(args.db)
        return

    if args.export:
        export_for_bidking(args.db)
        return

    config = AnalysisConfig(scraper_db=args.db, model=args.model)

    if args.phase1 or args.all:
        extractor = TextExtractor(config)
        extractor.run_extraction(args.limit)

    if args.phase2 or args.all:
        analyzer = OpportunityAnalyzer(config)
        await analyzer.run_analysis(args.limit)

    if not (args.phase1 or args.phase2 or args.all):
        print("Use --phase1, --phase2, or --all to run analysis")
        print("Use --stats to see current progress")
        print_stats(args.db)


if __name__ == "__main__":
    asyncio.run(main())
