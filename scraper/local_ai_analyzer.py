"""
Local AI Analysis System for SAM.gov Opportunities

Uses Ollama with a high-quality local LLM (Qwen2.5-32B) to analyze
PDF attachments and extract structured data matching BidKing's schema.

Prioritizes QUALITY over speed - uses larger model with careful prompting.
"""

import asyncio
import json
import logging
import re
import sqlite3
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Optional, Dict, List, Any
from concurrent.futures import ThreadPoolExecutor

import httpx

# PDF extraction
try:
    import pypdf
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

# DOCX extraction
try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# XLSX extraction
try:
    import openpyxl
    HAS_XLSX = True
except ImportError:
    HAS_XLSX = False

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('ai_analysis.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)


# =============================================================================
# Configuration
# =============================================================================

@dataclass
class AnalysisConfig:
    """Configuration for the analysis system."""

    # Ollama settings
    ollama_host: str = "http://localhost:11434"

    # Model selection - prioritizing quality
    # Options (in order of quality for your 4080):
    # - qwen2.5:32b-instruct-q4_K_M  (~18GB, best quality, some CPU offload)
    # - qwen2.5:14b-instruct         (~9GB, fits fully in VRAM, excellent quality)
    # - mistral-small:22b-instruct   (~12GB, fast with good quality)
    model: str = "qwen2.5:14b-instruct"  # Best balance of quality and speed for 4080

    # Fallback model if primary unavailable
    fallback_model: str = "qwen2.5:7b-instruct"

    # Database paths
    scraper_db: str = "bidking_sam.db"

    # PDF storage
    pdf_dir: str = "pdfs"

    # Processing settings
    max_text_length: int = 32000  # Max chars to send to LLM (quality over speed)
    batch_size: int = 1  # Process one at a time for quality

    # Timeouts (generous for quality)
    llm_timeout: int = 300  # 5 minutes per analysis
    download_timeout: int = 120


# =============================================================================
# Database Schema Extension
# =============================================================================

def init_analysis_tables(db_path: str):
    """Add AI analysis tables to the scraper database."""

    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()

    # AI Analysis results table
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS ai_analysis (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            opportunity_id TEXT NOT NULL,
            attachment_id INTEGER,

            -- Analysis status
            status TEXT DEFAULT 'pending',  -- pending, analyzing, completed, failed, skipped
            error_message TEXT,

            -- Extracted content
            text_content TEXT,
            text_length INTEGER,

            -- AI Analysis JSON (matches BidKing schema)
            ai_summary JSON,

            -- Metadata
            model_used TEXT,
            analysis_duration_seconds REAL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            analyzed_at TIMESTAMP,

            FOREIGN KEY (opportunity_id) REFERENCES opportunities(opportunity_id),
            FOREIGN KEY (attachment_id) REFERENCES attachments(id)
        )
    """)

    # Index for efficient queries
    cursor.execute("""
        CREATE INDEX IF NOT EXISTS idx_ai_analysis_status
        ON ai_analysis(status)
    """)

    cursor.execute("""
        CREATE INDEX IF NOT EXISTS idx_ai_analysis_opp_id
        ON ai_analysis(opportunity_id)
    """)

    # Track PDF download status - add columns if they don't exist
    # Get existing columns
    cursor.execute("PRAGMA table_info(attachments)")
    existing_cols = {row[1] for row in cursor.fetchall()}

    if 'pdf_downloaded' not in existing_cols:
        cursor.execute("ALTER TABLE attachments ADD COLUMN pdf_downloaded INTEGER DEFAULT 0")
    if 'pdf_local_path' not in existing_cols:
        cursor.execute("ALTER TABLE attachments ADD COLUMN pdf_local_path TEXT")
    if 'text_extracted' not in existing_cols:
        cursor.execute("ALTER TABLE attachments ADD COLUMN text_extracted INTEGER DEFAULT 0")
    if 'extracted_text' not in existing_cols:
        cursor.execute("ALTER TABLE attachments ADD COLUMN extracted_text TEXT")

    conn.commit()
    conn.close()
    logger.info(f"Initialized AI analysis tables in {db_path}")


# =============================================================================
# PDF Download & Text Extraction
# =============================================================================

class PDFProcessor:
    """Downloads PDFs and extracts text content."""

    def __init__(self, config: AnalysisConfig, proxy_file: Optional[str] = None):
        self.config = config
        self.pdf_dir = Path(config.pdf_dir)
        self.pdf_dir.mkdir(parents=True, exist_ok=True)

        # Load proxies if available
        self.proxies = []
        if proxy_file and Path(proxy_file).exists():
            with open(proxy_file) as f:
                for line in f:
                    line = line.strip()
                    if line and ':' in line:
                        parts = line.split(':')
                        if len(parts) >= 4:
                            host, port, user, passwd = parts[0], parts[1], parts[2], parts[3]
                            self.proxies.append(f"http://{user}:{passwd}@{host}:{port}")
            logger.info(f"Loaded {len(self.proxies)} proxies for PDF downloads")

        self.proxy_index = 0

    def _get_proxy(self) -> Optional[str]:
        """Get next proxy in rotation."""
        if not self.proxies:
            return None
        proxy = self.proxies[self.proxy_index % len(self.proxies)]
        self.proxy_index += 1
        return proxy

    async def download_pdf(
        self,
        url: str,
        opportunity_id: str,
        resource_id: str,
        filename: str
    ) -> Optional[Path]:
        """Download a PDF from SAM.gov."""

        # Create opportunity subdirectory
        opp_dir = self.pdf_dir / opportunity_id
        opp_dir.mkdir(exist_ok=True)

        # Sanitize filename
        safe_filename = re.sub(r'[<>:"/\\|?*]', '_', filename or resource_id)
        if not safe_filename.endswith('.pdf'):
            safe_filename += '.pdf'

        local_path = opp_dir / safe_filename

        # Skip if already downloaded
        if local_path.exists() and local_path.stat().st_size > 0:
            logger.debug(f"Already downloaded: {local_path}")
            return local_path

        # Try download with proxy rotation
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
            "Accept": "application/pdf,*/*",
        }

        for attempt in range(3):
            proxy = self._get_proxy()
            try:
                async with httpx.AsyncClient(
                    timeout=self.config.download_timeout,
                    follow_redirects=True,
                    proxy=proxy
                ) as client:
                    response = await client.get(url, headers=headers)

                    if response.status_code == 200:
                        content = response.content

                        # Verify it's actually a PDF
                        if content[:4] == b'%PDF' or len(content) > 1000:
                            with open(local_path, 'wb') as f:
                                f.write(content)
                            logger.info(f"Downloaded: {filename} ({len(content):,} bytes)")
                            return local_path
                        else:
                            logger.warning(f"Not a valid PDF: {filename}")
                            return None
                    else:
                        logger.warning(f"Download failed ({response.status_code}): {filename}")

            except Exception as e:
                logger.warning(f"Download attempt {attempt+1} failed: {e}")
                await asyncio.sleep(2)

        return None

    def extract_text(self, file_path: Path) -> Optional[str]:
        """Extract text from any supported document type (PDF, DOCX, XLSX, TXT, DOC)."""

        if not file_path.exists():
            return None

        suffix = file_path.suffix.lower()
        text = ""

        # Route to appropriate extractor based on file type
        if suffix == '.pdf':
            text = self._extract_pdf_text(file_path)
        elif suffix == '.docx':
            text = self._extract_docx_text(file_path)
        elif suffix == '.doc':
            text = self._extract_doc_text(file_path)
        elif suffix in ['.xlsx', '.xls']:
            text = self._extract_xlsx_text(file_path)
        elif suffix == '.txt':
            text = self._extract_txt_text(file_path)
        elif suffix == '.rtf':
            text = self._extract_txt_text(file_path)  # Basic RTF as text
        else:
            logger.warning(f"Unsupported file type: {suffix}")
            return None

        return text.strip() if text else None

    def _extract_pdf_text(self, pdf_path: Path) -> Optional[str]:
        """Extract text from a PDF file."""
        text = ""

        # Try pdfplumber first (better for tables and complex layouts)
        if HAS_PDFPLUMBER:
            try:
                with pdfplumber.open(pdf_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n\n"

                if text.strip():
                    logger.debug(f"Extracted {len(text)} chars with pdfplumber")
                    return text.strip()
            except Exception as e:
                logger.warning(f"pdfplumber failed: {e}")

        # Fallback to pypdf
        if HAS_PYPDF:
            try:
                reader = pypdf.PdfReader(pdf_path)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"

                if text.strip():
                    logger.debug(f"Extracted {len(text)} chars with pypdf")
                    return text.strip()
            except Exception as e:
                logger.warning(f"pypdf failed: {e}")

        return None

    def _extract_docx_text(self, docx_path: Path) -> Optional[str]:
        """Extract text from a DOCX file."""
        if not HAS_DOCX:
            logger.warning("python-docx not installed, cannot extract DOCX")
            return None

        try:
            doc = DocxDocument(docx_path)
            text_parts = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            text = "\n\n".join(text_parts)
            if text.strip():
                logger.debug(f"Extracted {len(text)} chars from DOCX")
                return text.strip()
        except Exception as e:
            logger.warning(f"DOCX extraction failed: {e}")

        return None

    def _extract_doc_text(self, doc_path: Path) -> Optional[str]:
        """Extract text from old .doc file using antiword or fallback."""
        import subprocess

        try:
            # Try antiword first (common on Linux)
            result = subprocess.run(
                ['antiword', str(doc_path)],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0 and result.stdout.strip():
                logger.debug(f"Extracted {len(result.stdout)} chars from DOC with antiword")
                return result.stdout.strip()
        except (subprocess.TimeoutExpired, FileNotFoundError):
            pass
        except Exception as e:
            logger.warning(f"antiword failed: {e}")

        # Fallback: try reading as text (sometimes works for simple docs)
        try:
            with open(doc_path, 'rb') as f:
                content = f.read()
                # Try to extract readable ASCII text
                text = ''.join(chr(b) for b in content if 32 <= b < 127 or b in (10, 13, 9))
                text = ' '.join(text.split())  # Normalize whitespace
                if len(text) > 100:
                    logger.debug(f"Extracted {len(text)} chars from DOC (raw)")
                    return text
        except Exception as e:
            logger.warning(f"DOC raw extraction failed: {e}")

        return None

    def _extract_xlsx_text(self, xlsx_path: Path) -> Optional[str]:
        """Extract text from XLSX/XLS file."""
        if not HAS_XLSX:
            logger.warning("openpyxl not installed, cannot extract XLSX")
            return None

        try:
            wb = openpyxl.load_workbook(xlsx_path, read_only=True, data_only=True)
            text_parts = []

            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                text_parts.append(f"=== Sheet: {sheet_name} ===")

                for row in sheet.iter_rows():
                    row_values = []
                    for cell in row:
                        if cell.value is not None:
                            row_values.append(str(cell.value).strip())
                    if row_values:
                        text_parts.append(" | ".join(row_values))

            wb.close()

            text = "\n".join(text_parts)
            if text.strip():
                logger.debug(f"Extracted {len(text)} chars from XLSX")
                return text.strip()
        except Exception as e:
            logger.warning(f"XLSX extraction failed: {e}")

        return None

    def _extract_txt_text(self, txt_path: Path) -> Optional[str]:
        """Extract text from TXT file."""
        try:
            # Try UTF-8 first
            with open(txt_path, 'r', encoding='utf-8') as f:
                text = f.read()
                if text.strip():
                    logger.debug(f"Extracted {len(text)} chars from TXT")
                    return text.strip()
        except UnicodeDecodeError:
            pass

        try:
            # Fallback to latin-1
            with open(txt_path, 'r', encoding='latin-1') as f:
                text = f.read()
                if text.strip():
                    logger.debug(f"Extracted {len(text)} chars from TXT (latin-1)")
                    return text.strip()
        except Exception as e:
            logger.warning(f"TXT extraction failed: {e}")

        return None


# =============================================================================
# Local LLM Analysis (Ollama)
# =============================================================================

class LocalLLMAnalyzer:
    """Analyzes document text using local Ollama model."""

    # Detailed prompt for high-quality extraction
    ANALYSIS_PROMPT = """You are an expert federal contracting analyst. Analyze this government solicitation document and extract key information in a structured JSON format.

IMPORTANT: Be thorough and accurate. Extract ONLY information that is explicitly stated in the document. If information is not present, use null.

Document text:
{document_text}

---

Extract the following information and return ONLY a valid JSON object (no markdown, no explanation):

{{
    "summary": "2-3 sentence plain English summary of what the government is looking for",

    "estimated_value": {{
        "low": <number or null>,
        "high": <number or null>,
        "basis": "Explanation of how you derived the estimate (e.g., 'Based on 3 FTEs at $150K for 2 years')"
    }},

    "period_of_performance": "e.g., '1 base year + 4 option years' or '5 years' or null if not stated",

    "contract_type": "e.g., 'Firm Fixed Price', 'Time & Materials', 'Cost Plus Fixed Fee', 'IDIQ', or null",

    "clearance_required": "e.g., 'None', 'Public Trust', 'Secret', 'Top Secret', 'TS/SCI', or null",

    "labor_categories": [
        {{
            "title": "Role title",
            "quantity": <number or null>,
            "level": "Junior/Mid/Senior/Lead/SME or null"
        }}
    ],

    "technologies": ["List of specific technologies, platforms, tools, languages mentioned"],

    "certifications_required": ["List of required certifications like CMMI, ISO, FedRAMP, etc."],

    "location": "e.g., 'Remote', 'Washington, DC', 'On-site at contractor facility', or description of work location",

    "incumbent": "Current contractor name if this is a recompete, or null",

    "key_dates": {{
        "proposal_due": "YYYY-MM-DD or null",
        "questions_due": "YYYY-MM-DD or null",
        "site_visit": "YYYY-MM-DD or null"
    }},

    "evaluation_factors": ["List of how proposals will be evaluated, e.g., 'Technical Approach', 'Past Performance', 'Price'"],

    "small_business_set_aside": "e.g., '8(a)', 'SDVOSB', 'HUBZone', 'Full and Open', or null",

    "naics_codes": ["List of NAICS codes mentioned"],

    "deliverables": ["Key deliverables or work products expected"]
}}

Return ONLY the JSON object, no other text."""

    def __init__(self, config: AnalysisConfig):
        self.config = config
        self.model = config.model

    async def check_ollama_available(self) -> bool:
        """Check if Ollama is running and the model is available."""
        try:
            async with httpx.AsyncClient(timeout=10) as client:
                response = await client.get(f"{self.config.ollama_host}/api/tags")
                if response.status_code == 200:
                    data = response.json()
                    models = [m['name'] for m in data.get('models', [])]

                    # Check for our preferred model
                    if any(self.model in m for m in models):
                        logger.info(f"Found model: {self.model}")
                        return True

                    # Check for fallback
                    if any(self.config.fallback_model in m for m in models):
                        logger.info(f"Using fallback model: {self.config.fallback_model}")
                        self.model = self.config.fallback_model
                        return True

                    logger.warning(f"Model not found. Available: {models}")
                    return False
        except Exception as e:
            logger.error(f"Ollama not available: {e}")
            return False

    async def analyze_document(self, text: str) -> Optional[Dict[str, Any]]:
        """Analyze document text and return structured extraction."""

        # Truncate if too long (but keep as much as possible for quality)
        if len(text) > self.config.max_text_length:
            # Keep beginning and end (usually most important parts)
            half = self.config.max_text_length // 2
            text = text[:half] + "\n\n[...content truncated...]\n\n" + text[-half:]

        prompt = self.ANALYSIS_PROMPT.format(document_text=text)

        start_time = datetime.now()

        try:
            async with httpx.AsyncClient(timeout=self.config.llm_timeout) as client:
                response = await client.post(
                    f"{self.config.ollama_host}/api/generate",
                    json={
                        "model": self.model,
                        "prompt": prompt,
                        "stream": False,
                        "options": {
                            "temperature": 0.1,  # Low temperature for accuracy
                            "num_predict": 4096,  # Allow long responses
                            "top_p": 0.9,
                        }
                    }
                )

                if response.status_code != 200:
                    logger.error(f"Ollama error: {response.status_code}")
                    return None

                data = response.json()
                response_text = data.get('response', '')

                # Parse JSON from response
                json_match = self._extract_json(response_text)
                if json_match:
                    duration = (datetime.now() - start_time).total_seconds()
                    logger.info(f"Analysis completed in {duration:.1f}s")
                    return json_match
                else:
                    logger.warning(f"Could not parse JSON from response")
                    return None

        except asyncio.TimeoutError:
            logger.error(f"Analysis timed out after {self.config.llm_timeout}s")
            return None
        except Exception as e:
            logger.error(f"Analysis error: {e}")
            return None

    def _extract_json(self, text: str) -> Optional[Dict]:
        """Extract JSON from LLM response."""

        # Try direct parse first
        try:
            return json.loads(text)
        except:
            pass

        # Remove markdown code blocks
        text = re.sub(r'```json\s*', '', text)
        text = re.sub(r'```\s*', '', text)

        # Try to find JSON object
        try:
            # Find first { and last }
            start = text.find('{')
            end = text.rfind('}')
            if start != -1 and end != -1:
                json_str = text[start:end+1]
                return json.loads(json_str)
        except:
            pass

        return None


# =============================================================================
# Main Analysis Pipeline
# =============================================================================

class AIAnalysisPipeline:
    """Complete pipeline for local AI analysis of SAM.gov opportunities."""

    def __init__(
        self,
        config: Optional[AnalysisConfig] = None,
        proxy_file: str = "/home/peteylinux/Downloads/Webshare 100 proxies(1).txt"
    ):
        self.config = config or AnalysisConfig()
        self.pdf_processor = PDFProcessor(self.config, proxy_file)
        self.llm_analyzer = LocalLLMAnalyzer(self.config)

        # PDF directory for checking local files
        self.pdf_dir = Path(self.config.pdf_dir)

        # Initialize database
        init_analysis_tables(self.config.scraper_db)

    def get_pending_attachments(self, limit: int = 100) -> List[Dict]:
        """Get PDF attachments that haven't been analyzed yet."""

        conn = sqlite3.connect(self.config.scraper_db)
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()

        cursor.execute("""
            SELECT
                a.id,
                a.opportunity_id,
                a.resource_id,
                a.filename,
                a.download_url,
                a.mime_type,
                a.access_level,
                a.pdf_local_path,
                o.title as opportunity_title
            FROM attachments a
            JOIN opportunities o ON a.opportunity_id = o.opportunity_id
            LEFT JOIN ai_analysis aa ON a.id = aa.attachment_id
            WHERE
                (
                    a.mime_type LIKE '%pdf%' OR a.filename LIKE '%.pdf'
                    OR a.mime_type LIKE '%word%' OR a.filename LIKE '%.docx' OR a.filename LIKE '%.doc'
                    OR a.mime_type LIKE '%excel%' OR a.mime_type LIKE '%spreadsheet%' OR a.filename LIKE '%.xlsx' OR a.filename LIKE '%.xls'
                    OR a.mime_type LIKE '%text%' OR a.filename LIKE '%.txt'
                    OR a.filename LIKE '%.rtf'
                )
                AND a.access_level = 'public'
                AND a.download_url IS NOT NULL
                AND aa.id IS NULL
            ORDER BY o.response_deadline ASC
            LIMIT ?
        """, (limit,))

        results = [dict(row) for row in cursor.fetchall()]
        conn.close()

        return results

    def save_analysis(
        self,
        opportunity_id: str,
        attachment_id: int,
        text_content: str,
        ai_summary: Optional[Dict],
        model_used: str,
        duration: float,
        status: str = "completed",
        error: Optional[str] = None
    ):
        """Save analysis results to database."""

        conn = sqlite3.connect(self.config.scraper_db)
        cursor = conn.cursor()

        cursor.execute("""
            INSERT INTO ai_analysis (
                opportunity_id, attachment_id, status, error_message,
                text_content, text_length, ai_summary,
                model_used, analysis_duration_seconds, analyzed_at
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            opportunity_id,
            attachment_id,
            status,
            error,
            text_content,
            len(text_content) if text_content else 0,
            json.dumps(ai_summary) if ai_summary else None,
            model_used,
            duration,
            datetime.now().isoformat() if status == "completed" else None
        ))

        conn.commit()
        conn.close()

    async def process_attachment(self, attachment: Dict) -> bool:
        """Process a single attachment through the full pipeline."""

        opp_id = attachment['opportunity_id']
        att_id = attachment['id']
        filename = attachment['filename'] or attachment['resource_id']
        url = attachment['download_url']

        logger.info(f"Processing: {filename} (opp: {opp_id[:8]}...)")

        start_time = datetime.now()

        # Step 1: Check for locally downloaded PDF first (from batch_pdf_downloader)
        pdf_path = None

        # Check if already downloaded by batch downloader
        local_pdf_path = attachment.get('pdf_local_path')
        if local_pdf_path and Path(local_pdf_path).exists():
            pdf_path = Path(local_pdf_path)
            logger.debug(f"Using pre-downloaded PDF: {pdf_path}")
        else:
            # Check in standard location (pdfs/opportunity_id/filename.pdf)
            safe_filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
            if not safe_filename.lower().endswith('.pdf'):
                safe_filename += '.pdf'
            potential_path = self.pdf_dir / opp_id / safe_filename

            if potential_path.exists() and potential_path.stat().st_size > 1000:
                pdf_path = potential_path
                logger.debug(f"Found local PDF at: {pdf_path}")

        # Step 2: Download PDF only if not found locally
        if not pdf_path:
            logger.debug(f"Downloading PDF (not found locally)...")
            pdf_path = await self.pdf_processor.download_pdf(
                url, opp_id, attachment['resource_id'], filename
            )

        if not pdf_path:
            self.save_analysis(
                opp_id, att_id, None, None,
                self.llm_analyzer.model, 0,
                status="failed", error="Download failed"
            )
            return False

        # Step 2: Extract text
        text = self.pdf_processor.extract_text(pdf_path)

        if not text or len(text) < 100:
            self.save_analysis(
                opp_id, att_id, text, None,
                self.llm_analyzer.model, 0,
                status="skipped", error="No extractable text"
            )
            return False

        # Step 3: AI Analysis
        analysis = await self.llm_analyzer.analyze_document(text)

        duration = (datetime.now() - start_time).total_seconds()

        if analysis:
            self.save_analysis(
                opp_id, att_id, text, analysis,
                self.llm_analyzer.model, duration,
                status="completed"
            )
            logger.info(f"✅ Completed analysis for {filename}")
            return True
        else:
            self.save_analysis(
                opp_id, att_id, text, None,
                self.llm_analyzer.model, duration,
                status="failed", error="LLM analysis failed"
            )
            return False

    async def run(self, limit: Optional[int] = None):
        """Run the analysis pipeline."""

        logger.info("=" * 60)
        logger.info("Local AI Analysis Pipeline")
        logger.info("=" * 60)

        # Check Ollama
        if not await self.llm_analyzer.check_ollama_available():
            logger.error("Ollama not available. Please run:")
            logger.error("  curl -fsSL https://ollama.com/install.sh | sudo sh")
            logger.error(f"  ollama pull {self.config.model}")
            return

        logger.info(f"Using model: {self.llm_analyzer.model}")

        # Get pending attachments
        attachments = self.get_pending_attachments(limit or 10000)
        logger.info(f"Found {len(attachments)} PDFs to analyze")

        if not attachments:
            logger.info("No pending attachments. Analysis complete!")
            return

        # Process one at a time for quality
        completed = 0
        failed = 0

        for i, attachment in enumerate(attachments, 1):
            logger.info(f"\n[{i}/{len(attachments)}] Processing...")

            try:
                success = await self.process_attachment(attachment)
                if success:
                    completed += 1
                else:
                    failed += 1
            except Exception as e:
                logger.error(f"Error processing attachment: {e}")
                failed += 1

            # Brief pause between analyses
            if i < len(attachments):
                await asyncio.sleep(1)

        logger.info("\n" + "=" * 60)
        logger.info(f"Analysis Complete!")
        logger.info(f"  Completed: {completed}")
        logger.info(f"  Failed: {failed}")
        logger.info("=" * 60)

    def get_stats(self) -> Dict:
        """Get analysis statistics."""

        conn = sqlite3.connect(self.config.scraper_db)
        cursor = conn.cursor()

        stats = {}

        # Supported file type filter
        file_filter = """
            (
                mime_type LIKE '%pdf%' OR filename LIKE '%.pdf'
                OR mime_type LIKE '%word%' OR filename LIKE '%.docx' OR filename LIKE '%.doc'
                OR mime_type LIKE '%excel%' OR mime_type LIKE '%spreadsheet%' OR filename LIKE '%.xlsx' OR filename LIKE '%.xls'
                OR mime_type LIKE '%text%' OR filename LIKE '%.txt'
                OR filename LIKE '%.rtf'
            )
        """

        # Total documents (all supported types)
        cursor.execute(f"""
            SELECT COUNT(*) FROM attachments
            WHERE {file_filter}
        """)
        stats['total_docs'] = cursor.fetchone()[0]

        # Analyzed
        cursor.execute("SELECT COUNT(*) FROM ai_analysis WHERE status = 'completed'")
        stats['analyzed'] = cursor.fetchone()[0]

        # Failed
        cursor.execute("SELECT COUNT(*) FROM ai_analysis WHERE status = 'failed'")
        stats['failed'] = cursor.fetchone()[0]

        # Pending
        cursor.execute(f"""
            SELECT COUNT(*) FROM attachments a
            LEFT JOIN ai_analysis aa ON a.id = aa.attachment_id
            WHERE {file_filter}
            AND aa.id IS NULL
        """)
        stats['pending'] = cursor.fetchone()[0]

        conn.close()
        return stats


# =============================================================================
# Export for BidKing
# =============================================================================

def export_for_bidking(
    db_path: str = "bidking_sam.db",
    output_file: str = "bidking_ai_import.json"
) -> int:
    """
    Export AI analysis results in BidKing-compatible format.

    Creates a JSON file that can be imported into BidKing's database.
    """

    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()

    # Get all completed analyses with opportunity data
    cursor.execute("""
        SELECT
            o.opportunity_id,
            o.solicitation_number,
            o.title,
            o.agency_name,
            o.response_deadline,
            aa.ai_summary,
            aa.text_content,
            aa.model_used,
            aa.analyzed_at,
            a.filename as source_document
        FROM ai_analysis aa
        JOIN opportunities o ON aa.opportunity_id = o.opportunity_id
        JOIN attachments a ON aa.attachment_id = a.id
        WHERE aa.status = 'completed' AND aa.ai_summary IS NOT NULL
        ORDER BY o.opportunity_id, aa.analyzed_at
    """)

    # Group by opportunity
    opportunities = {}

    for row in cursor:
        opp_id = row['opportunity_id']

        if opp_id not in opportunities:
            opportunities[opp_id] = {
                "opportunity_id": opp_id,
                "solicitation_number": row['solicitation_number'],
                "title": row['title'],
                "agency_name": row['agency_name'],
                "response_deadline": row['response_deadline'],
                "ai_analyses": [],
                "source_documents": []
            }

        ai_summary = json.loads(row['ai_summary']) if row['ai_summary'] else {}

        opportunities[opp_id]['ai_analyses'].append({
            "summary": ai_summary,
            "model_used": row['model_used'],
            "analyzed_at": row['analyzed_at']
        })

        if row['source_document']:
            opportunities[opp_id]['source_documents'].append(row['source_document'])

    conn.close()

    # Merge multiple analyses per opportunity
    export_data = []

    for opp_id, opp in opportunities.items():
        # Merge AI summaries from multiple documents
        merged_summary = _merge_summaries([a['summary'] for a in opp['ai_analyses']])

        export_data.append({
            "opportunity_id": opp['opportunity_id'],
            "solicitation_number": opp['solicitation_number'],
            "title": opp['title'],
            "agency_name": opp['agency_name'],
            "response_deadline": opp['response_deadline'],
            "ai_summary": merged_summary,
            "source_documents": list(set(opp['source_documents'])),
            "analysis_metadata": {
                "model_used": opp['ai_analyses'][0]['model_used'],
                "analyzed_at": opp['ai_analyses'][-1]['analyzed_at'],
                "num_documents_analyzed": len(opp['ai_analyses'])
            }
        })

    # Write export file
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump({
            "exported_at": datetime.now().isoformat(),
            "total_opportunities": len(export_data),
            "opportunities": export_data
        }, f, indent=2)

    logger.info(f"Exported {len(export_data)} opportunities to {output_file}")
    return len(export_data)


def _merge_summaries(summaries: List[Dict]) -> Dict:
    """Merge multiple AI summaries into one comprehensive summary."""

    if not summaries:
        return {}

    if len(summaries) == 1:
        return summaries[0]

    # Start with first summary
    merged = summaries[0].copy()

    # Merge additional summaries
    for summary in summaries[1:]:
        # Concatenate summaries
        if summary.get('summary') and merged.get('summary'):
            if summary['summary'] not in merged['summary']:
                merged['summary'] = merged['summary'] + " " + summary['summary']

        # Take best estimated value
        if summary.get('estimated_value', {}).get('low'):
            if not merged.get('estimated_value', {}).get('low'):
                merged['estimated_value'] = summary['estimated_value']

        # Merge lists (dedup)
        for list_field in ['technologies', 'certifications_required', 'labor_categories',
                          'evaluation_factors', 'naics_codes', 'deliverables']:
            if summary.get(list_field):
                existing = merged.get(list_field, [])
                for item in summary[list_field]:
                    if item not in existing:
                        existing.append(item)
                merged[list_field] = existing

        # Take non-null values
        for field in ['period_of_performance', 'contract_type', 'clearance_required',
                     'location', 'incumbent', 'small_business_set_aside']:
            if summary.get(field) and not merged.get(field):
                merged[field] = summary[field]

    return merged


# =============================================================================
# CLI
# =============================================================================

async def main():
    import argparse

    parser = argparse.ArgumentParser(description="Local AI Analysis for SAM.gov")
    parser.add_argument("--db", default="bidking_sam.db", help="Database path")
    parser.add_argument("--limit", type=int, help="Max attachments to process")
    parser.add_argument("--model", default="qwen2.5:14b-instruct", help="Ollama model")
    parser.add_argument("--export", action="store_true", help="Export results for BidKing")
    parser.add_argument("--stats", action="store_true", help="Show analysis stats")

    args = parser.parse_args()

    config = AnalysisConfig(
        scraper_db=args.db,
        model=args.model
    )

    pipeline = AIAnalysisPipeline(config)

    if args.stats:
        stats = pipeline.get_stats()
        print("\n" + "=" * 40)
        print("AI Analysis Statistics")
        print("=" * 40)
        print(f"Total Docs:  {stats['total_docs']:,}")
        print(f"Analyzed:    {stats['analyzed']:,}")
        print(f"Failed:      {stats['failed']:,}")
        print(f"Pending:     {stats['pending']:,}")
        print("=" * 40)
        return

    if args.export:
        export_for_bidking(args.db)
        return

    await pipeline.run(args.limit)


if __name__ == "__main__":
    asyncio.run(main())
